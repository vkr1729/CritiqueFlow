import logging
import sys
from pathlib import Path, PurePath

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".md", ".txt", ".xlsx", ".xls", ".docx", ".pdf"}


def list_files(folder_path: str) -> list[str]:
    path = Path(folder_path)
    if not path.is_absolute():
        raise ValueError(f"folder_path must be absolute: {folder_path}")
    if not path.exists() or not path.is_dir():
        raise FileNotFoundError(f"Directory not found: {folder_path}")
    files = []
    for entry in sorted(path.iterdir()):
        if entry.is_file() and entry.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(entry.name)
    return files


def _validate_filename(filename: str) -> str:
    if "/" in filename or "\\" in filename:
        raise ValueError(f"Directory separators not allowed in filename: {filename}")
    parts = PurePath(filename).parts
    if ".." in parts:
        raise ValueError(f"Path traversal attempt: {filename}")
    return filename


def _read_text(path: Path, char_cap: int) -> str:
    with open(path, "r", encoding="utf-8") as f:
        content = f.read() if char_cap >= sys.maxsize else f.read(char_cap)
    return content


def _read_xlsx(path: Path, row_cap: int) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    lines = []
    row_count = 0
    for row in ws.iter_rows(values_only=True):
        if row_count >= row_cap:
            break
        lines.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
        row_count += 1
    wb.close()
    return "\n".join(lines)


def _read_docx(path: Path, char_cap: int) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    total = 0
    for para in doc.paragraphs:
        if total >= char_cap:
            break
        text = para.text
        parts.append(text)
        total += len(text)
    for table in doc.tables:
        if total >= char_cap:
            break
        for row in table.rows:
            if total >= char_cap:
                break
            row_text = " | ".join(cell.text for cell in row.cells)
            parts.append(row_text)
            total += len(row_text)
    return "\n".join(parts)[:char_cap]


def _read_pdf(path: Path, char_cap: int) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(str(path))
    parts = []
    total = 0
    for page in reader.pages:
        if total >= char_cap:
            break
        text = page.extract_text()
        if text:
            parts.append(text)
            total += len(text)
    return "\n".join(parts)[:char_cap]


def read_file(folder_path: str, filename: str) -> dict:
    from config.settings import settings

    filename = _validate_filename(filename)
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file extension: {ext}")

    folder = Path(folder_path)
    full_path = (folder / filename).resolve()
    if not full_path.is_file():
        raise FileNotFoundError(f"File not found: {full_path}")

    # Path traversal check — use relative_to() which is immune to prefix attacks
    folder_real = folder.resolve()
    try:
        full_path.relative_to(folder_real)
    except ValueError:
        raise ValueError(f"Path traversal blocked: {filename}")

    char_cap = settings.get_effective_file_char_cap()
    row_cap = settings.get_effective_file_row_cap()

    if ext in (".md", ".txt"):
        content = _read_text(full_path, char_cap)
    elif ext in (".xlsx", ".xls"):
        content = _read_xlsx(full_path, row_cap)
    elif ext == ".docx":
        content = _read_docx(full_path, char_cap)
    elif ext == ".pdf":
        content = _read_pdf(full_path, char_cap)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

    if len(content) > 100_000:
        logger.warning("File '%s' is %d characters — this may consume significant token quota.", filename, len(content))

    return {"filename": filename, "content": content}
